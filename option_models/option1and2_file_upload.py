# pip install python-docx 
# pip install --upgrade google-api-python-client google-auth-httplib2 google-auth-oauthlib
import pickle
import pandas as pd
from docx import Document
from Google import Create_Service
from googleapiclient.errors import HttpError
import warnings
import sys
warnings.filterwarnings('ignore')

CLIENT_SECRET_FILE = 'client_secret.json'
API_NAME = 'forms'
API_VERSION = 'v1'
SCOPES = ["https://www.googleapis.com/auth/forms.body", "https://www.googleapis.com/auth/forms.responses.readonly"]

try:
    service = pickle.load(open('service.pkl', 'rb'))
    service
except:
    service = Create_Service(CLIENT_SECRET_FILE, API_NAME, API_VERSION, SCOPES)
    with open('service.pkl', 'wb' ) as f:
        pickle.dump(service, f)
        f.close()

class Question:
    
    input_q = None 
    
    def __init__(self, q_id, text, type_q):
            self.q_id = q_id
            self.text = text
            self.type_q = type_q


def doc2df(filepath):
    
    doc = Document(filepath)

    q_dict = {}
    q_num = 1

    for i, para in enumerate(doc.paragraphs):

        if para.text[0] == '~':

            temp_q = Question(q_id=f'Q{q_num}', text=para.text.replace('~', ''), type_q=doc.paragraphs[i+1].text)

            try:
                if doc.paragraphs[i+2].text[0] != '~':
                    temp_q.input_q = doc.paragraphs[i+2].text
                else:
                    temp_q.input_q = None

            except IndexError:
                temp_q.input_q = None
                
            q_dict[f'q{q_num}'] = temp_q
        
            q_num += 1
        
        else:
            pass
            
    df = pd.DataFrame([q_dict[q].__dict__ for q in q_dict])
    
    return df

def import_data(filepath):
    if '.xlsx' in filepath:
        temp_df = pd.read_excel(filepath)
        data = pd.concat([pd.Series([f'Q{index+1}' for index in temp_df.index], name='q_id'), temp_df], axis=1)
        
    elif '.doc' in filepath:
        data = doc2df(filepath=filepath)
    
    else:
        print('Invalid File Format!!! : Only Excel and Word Document Files are Accepted.')
        
    return data

def create_question(index, title, type_q, required=True, options=None):
    
    temp_question = {
            "createItem": {
                "item": {
                    "title": title,
                    "questionItem": {
                        "question": {
                        }
                    },
                },
                "location": {
                    "index": index
                }
            }
        }

    if type_q.lower() in ['string', 'number']:
        temp_question["createItem"]['item']['questionItem']['question'] = {'required': required, 'textQuestion': {'paragraph':False}}
        
    elif type_q.lower() == 'text':
        temp_question["createItem"]['item']['questionItem']['question'] = {'required': required, 'textQuestion': {'paragraph':True}}
    
    elif type_q.lower() == 'choice':
        temp_question["createItem"]['item']['questionItem']['question'] = {'required': required, 'choiceQuestion': {'type': 'RADIO', 'options': options, 'shuffle': False}}
        
    elif type_q.lower() == 'multiple':
        temp_question["createItem"]['item']['questionItem']['question'] = {'required': required, 'choiceQuestion': {'type': 'CHECKBOX', 'options': options, 'shuffle': False}}
        
    elif type_q.lower() == 'date':
        temp_question["createItem"]['item']['questionItem']['question'] = {'required': required, 'dateQuestion': {"includeTime": False, "includeYear": True}}
    
    elif type_q.lower() == 'datetime':
        temp_question["createItem"]['item']['questionItem']['question'] = {'required': required, 'dateQuestion': {"includeTime": True, "includeYear": False}}
    
    elif type_q.lower() == 'time':
        temp_question["createItem"]['item']['questionItem']['question'] = {'required': required, 'timeQuestion': {"duration": False}}
        
    elif type_q.lower() == 'scale':
        temp_question["createItem"]['item']['questionItem']['question'] = {'required': required, 'scaleQuestion': {"high": 0, "highLabel": "High", "low": 5, "lowLabel": "Low"}}
        
    elif type_q.lower() == 'file':
        temp_question["createItem"]['item']['questionItem']['question'] = {'required': required, "fileUploadQuestion": {"maxFileSize": 1, "maxFiles": 10}}
        
    else:
        pass
    
    
        
    return temp_question

def create_form(filepath, title, service, description=None):
    
    data = import_data(filepath=filepath)
    
    NEW_FORM = {
        "info": {
            "title": title,
            "description": description
        }
    }
    
    q_list = []
    
    for i in data.index:
        if pd.notna(data.loc[i, 'input_q']):
            options = [{'value':ele.strip()} for ele in data.loc[i, 'input_q'].split(',')]
        else:
            options=None
        
        required = True if '*' in data.loc[i, 'text'] else False
        
        question = create_question(index=i, title=data.loc[i, 'text'], type_q=data.loc[i, 'type_q'], required=required, options=options)
        
        q_list.append(question)
        
    NEW_QUESTION = {"requests": q_list}
    
    # Creates the initial form
    result = service.forms().create(body=NEW_FORM).execute()
    # print('Form has been Created.')

    # Adds the question to the form
    question_setting = service.forms().batchUpdate(formId=result["formId"], body=NEW_QUESTION).execute()
    # print('Questions have been Added.')

    # Prints the result to show the question has been added
    get_result = service.forms().get(formId=result["formId"]).execute()

    # print('Visit the Responder URL:',get_result['responderUri'])
    print(str(get_result['responderUri']) + '|' + str(len(q_list)))
    
    return get_result['formId'], get_result['responderUri']

# Usage
file_path = sys.argv[1]
file_title = sys.argv[2]
create_form(filepath=file_path, service=service, title=file_title)
# create_form(filepath='sample.xlsx', service=service, title='Sample Test 1')